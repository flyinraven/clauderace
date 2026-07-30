"""Document extraction: text, images and figure captions.

Handles PDF (PyMuPDF), Word, plain text and JSON. The interesting work is in
separating genuine clinical figures from document chrome: the RANZCO written
reports carry a letterhead banner, and the OSCE reports are PowerPoint exports
in which the slide background is embedded as an image on all 101 pages.
"""

from __future__ import annotations

import hashlib
import io
import json
import logging
import re
from collections import Counter
from dataclasses import dataclass, field
from typing import Any

logger = logging.getLogger(__name__)

# How often to empty MuPDF's cache while walking a document. Measured on the
# 160-page 2024 Semester 1 paper: extraction left the process at 429 MB, of
# which 258 MB was that cache. The instance it runs on has 512 MB, and the job
# worker is a thread inside the API, so an ingest that overshoots takes the
# whole site down with it - which is what the memory alerts were.
SHRINK_EVERY_PAGES = 20


def release_reader_memory() -> None:
    """Give back what the PDF reader is holding but no longer needs.

    Two steps, because they free different things. MuPDF keeps decoded page
    content in a store of its own that `doc.close()` does not empty; shrinking
    it is what recovers the 258 MB. `malloc_trim` then returns the freed heap
    to the operating system rather than leaving it in the allocator's arenas,
    which is what makes the difference visible to a container memory limit.

    Both are best-effort: neither is available everywhere, and neither failing
    changes the result of the extraction.
    """
    try:
        import pymupdf

        pymupdf.TOOLS.store_shrink(100)
    except Exception:  # noqa: BLE001 - a cache that will not shrink is not fatal
        logger.debug("Could not shrink the MuPDF store")

    try:
        import ctypes

        ctypes.CDLL("libc.so.6").malloc_trim(0)
    except Exception:  # noqa: BLE001 - glibc only; a no-op on Windows and musl
        pass


# --- Decorative-image heuristics -----------------------------------------
# An image repeated on more than this many pages is a template, not a figure.
MAX_PAGES_FOR_FIGURE = 2
# Logos and rules are small in at least one dimension.
MIN_DIMENSION_PX = 100
# Banners are extremely wide relative to their height.
MAX_ASPECT_RATIO = 4.0
MIN_SIZE_BYTES = 4 * 1024
# Anything sitting entirely inside these bands of the page is running header or
# footer - a college crest, a page rule, an examiner-report watermark. A
# clinical photograph is never printed in the top eighth of the page with text
# flowing under it, and this costs nothing to check, where letting it through
# costs a vision call to reach the same conclusion.
HEADER_BAND = 0.12
FOOTER_BAND = 0.92

FIGURE_CAPTION_RE = re.compile(
    r"^\s*(?:figure|fig\.?|image)\s*([0-9]+[a-z]?)\s*[:.\-]?\s*(.*)$",
    re.IGNORECASE,
)

# PDF text layers carry the typesetter's ligature glyphs, so "a useful negative
# finding" extracts as "ﬁnding" (U+FB01). It renders acceptably and then breaks
# every search, and it reaches the grading model as a word it has to guess at.
# Targeted rather than a blanket NFKC pass, which would also rewrite the micro
# sign, fraction slashes and superscripts that clinical text uses deliberately.
TEXT_SUBSTITUTIONS = {
    "ﬀ": "ff", "ﬁ": "fi", "ﬂ": "fl", "ﬃ": "ffi",
    "ﬄ": "ffl", "ﬅ": "st", "ﬆ": "st",
    "Ĳ": "IJ", "ĳ": "ij", "Œ": "OE", "œ": "oe",
    # Smart punctuation, for the same searchability reason.
    "‘": "'", "’": "'", "‚": "'", "‛": "'",
    "“": '"', "”": '"', "„": '"', "‟": '"',
    "′": "'", "″": '"',
    # Separators that survive a copy-paste and then break word splitting.
    " ": " ", " ": " ", " ": " ", "​": "", "﻿": "",
    "­": "",
}
_SUBSTITUTION_RE = re.compile("|".join(map(re.escape, TEXT_SUBSTITUTIONS)))


def normalise_extracted_text(text: str) -> str:
    """Fold typesetter glyphs back into the characters people search for."""
    if not text:
        return text
    return _SUBSTITUTION_RE.sub(lambda m: TEXT_SUBSTITUTIONS[m.group(0)], text)


@dataclass
class ExtractedImage:
    data: bytes
    content_type: str
    width: int
    height: int
    page_number: int
    sha256: str
    # PDF user-space box (x0, y0, x1, y1); y grows downwards.
    bbox: tuple[float, float, float, float] | None = None
    label: str | None = None
    caption: str | None = None

    @property
    def aspect_ratio(self) -> float:
        if not self.width or not self.height:
            return 1.0
        return max(self.width, self.height) / max(1, min(self.width, self.height))


@dataclass
class ExtractedPage:
    number: int
    text: str
    images: list[ExtractedImage] = field(default_factory=list)
    # Page height in PDF user space, for judging what is header and footer.
    height: float = 0.0

    def __post_init__(self) -> None:
        # Every extractor funnels through here, so this is the one place the
        # normalisation has to happen.
        self.text = normalise_extracted_text(self.text)


@dataclass
class ExtractedDocument:
    pages: list[ExtractedPage]
    page_count: int
    metadata: dict[str, Any] = field(default_factory=dict)
    discarded_images: int = 0

    @property
    def full_text(self) -> str:
        return "\n".join(p.text for p in self.pages)

    @property
    def images(self) -> list[ExtractedImage]:
        return [img for page in self.pages for img in page.images]


def sha256_bytes(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


# --- Dispatch -------------------------------------------------------------
def extract_document(data: bytes, filename: str, content_type: str) -> ExtractedDocument:
    name = (filename or "").lower()
    if name.endswith(".pdf") or "pdf" in (content_type or ""):
        return extract_pdf(data)
    if name.endswith(".docx") or "wordprocessingml" in (content_type or ""):
        return extract_docx(data)
    if name.endswith(".json") or "json" in (content_type or ""):
        return extract_json(data)
    return extract_text(data)


# --- PDF ------------------------------------------------------------------
def extract_pdf(data: bytes) -> ExtractedDocument:
    import pymupdf

    doc = pymupdf.open(stream=data, filetype="pdf")
    try:
        raw_pages: list[ExtractedPage] = []
        hash_pages: dict[str, set[int]] = {}

        for index in range(len(doc)):
            page = doc[index]
            page_number = index + 1
            text = page.get_text() or ""

            # Map xref -> on-page rectangle so captions can be matched by
            # position and figures ordered top-to-bottom.
            placements: dict[int, tuple[float, float, float, float]] = {}
            try:
                for info in page.get_image_info(xrefs=True):
                    xref = info.get("xref")
                    if xref:
                        placements[xref] = tuple(info["bbox"])
            except Exception:  # noqa: BLE001 - placement is a nicety, not required
                logger.debug("No image placement info for page %s", page_number)

            images: list[ExtractedImage] = []
            for entry in page.get_images(full=True):
                xref = entry[0]
                try:
                    info = doc.extract_image(xref)
                except Exception:  # noqa: BLE001 - skip unreadable image objects
                    logger.debug("Could not extract image xref %s", xref)
                    continue
                blob = info["image"]
                digest = sha256_bytes(blob)
                hash_pages.setdefault(digest, set()).add(page_number)
                images.append(
                    ExtractedImage(
                        data=blob,
                        content_type=f"image/{'jpeg' if info['ext'] == 'jpg' else info['ext']}",
                        width=info.get("width", 0),
                        height=info.get("height", 0),
                        page_number=page_number,
                        sha256=digest,
                        bbox=placements.get(xref),
                    )
                )

            images.sort(key=lambda i: (i.bbox[1] if i.bbox else 0, i.bbox[0] if i.bbox else 0))
            raw_pages.append(
                ExtractedPage(
                    number=page_number, text=text, images=images,
                    height=float(page.rect.height or 0),
                )
            )

            # Everything wanted from this page has been copied out, so MuPDF's
            # decoded copy of it is dead weight from here on. Emptying the store
            # periodically keeps the peak flat across a long document instead of
            # letting it climb with the page count.
            if page_number % SHRINK_EVERY_PAGES == 0:
                release_reader_memory()

        kept_pages, discarded = _filter_decorative(raw_pages, hash_pages, len(doc))
        for page, source in zip(kept_pages, doc, strict=False):
            _attach_captions(page, source)

        return ExtractedDocument(
            pages=kept_pages,
            page_count=len(doc),
            metadata={k: v for k, v in (doc.metadata or {}).items() if v},
            discarded_images=discarded,
        )
    finally:
        doc.close()
        # `close()` releases the document, not the store behind it.
        release_reader_memory()


def _is_page_furniture(img: ExtractedImage, page_height: float) -> bool:
    """Whether this image lies wholly in the running header or footer."""
    if not img.bbox or page_height <= 0:
        return False
    top, bottom = img.bbox[1], img.bbox[3]
    return bottom <= page_height * HEADER_BAND or top >= page_height * FOOTER_BAND


def _filter_decorative(
    pages: list[ExtractedPage], hash_pages: dict[str, set[int]], page_count: int
) -> tuple[list[ExtractedPage], int]:
    """Drop letterheads, logos, slide backgrounds and other chrome.

    Verified against the supplied reports: the 2026 OSCE export yields 119 image
    objects but only 4 distinct ones, all template - this reduces it to zero
    figures, which is correct. The 2026 written report keeps the 10 genuine
    clinical figures and drops the 2 cover banners.
    """
    discarded = 0
    seen_hashes: Counter[str] = Counter()

    for page in pages:
        keep: list[ExtractedImage] = []
        for img in page.images:
            reason = None
            if len(hash_pages.get(img.sha256, set())) > MAX_PAGES_FOR_FIGURE:
                reason = "appears on multiple pages (template)"
            elif page.number == 1 and page_count > 2:
                reason = "cover page"
            elif min(img.width, img.height) < MIN_DIMENSION_PX:
                reason = "too small (logo)"
            elif img.aspect_ratio > MAX_ASPECT_RATIO:
                reason = "banner aspect ratio"
            elif len(img.data) < MIN_SIZE_BYTES:
                reason = "too few bytes"
            elif _is_page_furniture(img, page.height):
                reason = "sits in the header or footer band"
            elif seen_hashes[img.sha256]:
                reason = "duplicate"

            if reason:
                discarded += 1
                logger.debug("Discarding image on p%s: %s", page.number, reason)
                continue
            seen_hashes[img.sha256] += 1
            keep.append(img)
        page.images = keep

    return pages, discarded


def _attach_captions(page: ExtractedPage, source) -> None:
    """Associate "Figure N: ..." lines with the image they label.

    In the RANZCO reports the caption sits immediately above its figure, so each
    caption claims the nearest image whose top edge is below it.
    """
    if not page.images:
        return
    try:
        blocks = source.get_text("blocks") or []
    except Exception:  # noqa: BLE001
        return

    captions: list[tuple[float, str, str]] = []
    for block in blocks:
        if len(block) < 5:
            continue
        # Captions come straight from the PDF, not via ExtractedPage.
        y0, text = block[1], normalise_extracted_text(block[4] or "")
        for line in text.splitlines():
            match = FIGURE_CAPTION_RE.match(line.strip())
            if match:
                captions.append((y0, f"Figure {match.group(1)}", match.group(2).strip()))

    unclaimed = list(page.images)
    for y0, label, caption in captions:
        below = [img for img in unclaimed if img.bbox and img.bbox[1] >= y0 - 5]
        target = min(below, key=lambda i: i.bbox[1]) if below else (unclaimed[0] if unclaimed else None)
        if target is None:
            break
        target.label = label
        target.caption = caption or None
        unclaimed.remove(target)


# --- Word / text / JSON ---------------------------------------------------
def extract_docx(data: bytes) -> ExtractedDocument:
    from docx import Document

    document = Document(io.BytesIO(data))

    lines: list[str] = []
    for para in document.paragraphs:
        lines.append(para.text)
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                lines.append(" | ".join(cells))

    images: list[ExtractedImage] = []
    seen: set[str] = set()
    for rel in document.part.rels.values():
        if "image" not in rel.reltype:
            continue
        try:
            blob = rel.target_part.blob
        except Exception:  # noqa: BLE001
            continue
        digest = sha256_bytes(blob)
        if digest in seen or len(blob) < MIN_SIZE_BYTES:
            continue
        seen.add(digest)
        width, height = _image_dimensions(blob)
        if min(width, height) < MIN_DIMENSION_PX:
            continue
        images.append(
            ExtractedImage(
                data=blob,
                content_type=_content_type_for(rel.target_part.partname),
                width=width,
                height=height,
                page_number=1,
                sha256=digest,
            )
        )

    text = "\n".join(lines)
    return ExtractedDocument(
        pages=[ExtractedPage(number=1, text=text, images=images)], page_count=1
    )


def extract_text(data: bytes) -> ExtractedDocument:
    for encoding in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):
        try:
            text = data.decode(encoding)
            break
        except UnicodeDecodeError:
            continue
    else:
        text = data.decode("utf-8", errors="replace")
    return ExtractedDocument(pages=[ExtractedPage(number=1, text=text)], page_count=1)


def extract_json(data: bytes) -> ExtractedDocument:
    """Pass structured JSON through as pretty-printed text.

    A JSON upload that already matches the question schema bypasses the model
    entirely - see `app.services.ingest.structure.parse_prestructured`.
    """
    doc = extract_text(data)
    try:
        payload = json.loads(doc.full_text)
    except json.JSONDecodeError as exc:
        raise ValueError(f"File is not valid JSON: {exc}") from exc
    doc.metadata["json"] = payload
    doc.pages[0].text = json.dumps(payload, indent=2, ensure_ascii=False)
    return doc


def _image_dimensions(blob: bytes) -> tuple[int, int]:
    try:
        from PIL import Image as PILImage

        with PILImage.open(io.BytesIO(blob)) as im:
            return im.width, im.height
    except Exception:  # noqa: BLE001 - dimensions are advisory
        return 0, 0


def _content_type_for(partname: Any) -> str:
    ext = str(partname).rsplit(".", 1)[-1].lower()
    return {"jpg": "image/jpeg", "jpeg": "image/jpeg", "png": "image/png",
            "gif": "image/gif", "bmp": "image/bmp", "tiff": "image/tiff",
            "webp": "image/webp"}.get(ext, "image/png")
